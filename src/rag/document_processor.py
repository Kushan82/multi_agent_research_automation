import os
import hashlib
from typing import List, Dict, Any
from pathlib import Path
import tiktoken
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document
import PyPDF2
import docx
from bs4 import BeautifulSoup
import requests
from utils.logger import setup_logger

logger = setup_logger("DocumentProcessor")

class DocumentProcessor:
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", ". ", " ", ""]
        )
        self.encoding = tiktoken.get_encoding("cl100k_base")
    
    def process_file(self, file_path: str) -> List[Document]:
        """Process a single file and return chunks"""
        try:
            extension = Path(file_path).suffix.lower()
            
            if extension == '.pdf':
                text = self._extract_pdf_text(file_path)
            elif extension == '.docx':
                text = self._extract_docx_text(file_path)
            elif extension == '.txt':
                text = self._extract_txt_text(file_path)
            else:
                raise ValueError(f"Unsupported file type: {extension}")
            
            return self._chunk_text(text, file_path)
            
        except Exception as e:
            logger.error(f"Failed to process {file_path}: {str(e)}")
            return []
    
    def process_url(self, url: str) -> List[Document]:
        """Process web content and return chunks"""
        try:
            response = requests.get(url, timeout=10)
            response.raise_for_status()
            
            soup = BeautifulSoup(response.content, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            text = soup.get_text()
            return self._chunk_text(text, url)
            
        except Exception as e:
            logger.error(f"Failed to process URL {url}: {str(e)}")
            return []
    
    def _extract_pdf_text(self, file_path: str) -> str:
        """Extract text from PDF file"""
        text = ""
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
        return text
    
    def _extract_docx_text(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        doc = docx.Document(file_path)
        return "\n".join([paragraph.text for paragraph in doc.paragraphs])
    
    def _extract_txt_text(self, file_path: str) -> str:
        """Extract text from TXT file"""
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    
    def _chunk_text(self, text: str, source: str) -> List[Document]:
        """Split text into chunks"""
        # Clean the text
        text = self._clean_text(text)
        
        # Split into chunks
        chunks = self.text_splitter.split_text(text)
        
        documents = []
        for i, chunk in enumerate(chunks):
            doc = Document(
                page_content=chunk,
                metadata={
                    "source": source,
                    "chunk_id": i,
                    "chunk_size": len(chunk),
                    "token_count": len(self.encoding.encode(chunk)),
                    "content_hash": hashlib.md5(chunk.encode()).hexdigest()
                }
            )
            documents.append(doc)
        
        return documents
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize text"""
        # Remove extra whitespace
        text = " ".join(text.split())
        
        # Remove non-printable characters
        text = ''.join(char for char in text if char.isprintable() or char.isspace())
        
        return text
    
    def get_token_count(self, text: str) -> int:
        """Get token count for text"""
        return len(self.encoding.encode(text))